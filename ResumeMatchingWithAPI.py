import os
import boto3
import psycopg2
from fastapi import HTTPException
import sshtunnel
import re
import pathlib
from typing import List, Optional, Dict, Set, Any, Union, Tuple
import fitz  # PyMuPDF for PDF extraction
from docx import Document
from datetime import datetime
import openai
import logging
import traceback
from dotenv import load_dotenv

load_dotenv()
# -------------------------- Database Connection --------------------------





# -------------------------- S3 File Download --------------------------


class S3FileDownload:

    @staticmethod
    def parse_s3_uri(s3_uri: str):
        """
        Parses an S3 URI to extract the bucket name and key.
        Example: s3://bucket-name/folder/file.pdf -> ('bucket-name', 'folder/file.pdf')
        """
        if s3_uri.startswith("s3://"):
            s3_uri = s3_uri[5:]
        bucket_name, key = s3_uri.split("/", 1)
        return bucket_name, key

    @staticmethod
    def download_s3_file(s3_uri: str, local_path: str):
        """
        Downloads a file from S3 using given credentials
        and stores it locally at local_path.
        """
        aws_access_key_id = os.getenv("AWS_ACCESS_KEY_ID")
        aws_secret_access_key = os.getenv("AWS_SECRET_ACCESS_KEY")

        bucket_name, key = S3FileDownload.parse_s3_uri(s3_uri)

        s3_client = boto3.client(
            "s3",
            aws_access_key_id=aws_access_key_id,
            aws_secret_access_key=aws_secret_access_key
        )

        print(f"[INFO] Downloading from S3: bucket={bucket_name}, key={key}")
        s3_client.download_file(bucket_name, key, local_path)
        print(f"[INFO] Download complete -> {local_path}")


#  -------------------------- Access JOBS Table --------------------------
class JobDataAccess:
    @staticmethod
    def get_job_details(job_id: int) -> dict:
        """
        Fetches the job details including description, title, and years of experience for a specific job_id.

        Args:
            job_id (int): The ID of the job

        Returns:
            dict: Dictionary containing job details

        Raises:
            HTTPException: If job_id not found or database error occurs
        """
        with DatabaseConnection() as db_conn:
            try:
                query = """
                    SELECT job_desc, job_title, years_of_exp
                    FROM public.jobs
                    WHERE job_id = %s
                """
                db_conn.cur.execute(query, (job_id,))

                result = db_conn.cur.fetchone()

                if not result:
                    raise HTTPException(
                        status_code=404,
                        detail=f"Job details for job ID {job_id} not found.",
                    )

                return {
                    "job_desc": result[0],
                    "job_title": result[1],
                    "years_of_exp": result[2]
                }

            except HTTPException:
                raise
            except Exception as e:
                raise HTTPException(
                    status_code=500,
                    detail=f"Database error while fetching job details: {str(e)}",
                )
    
    @staticmethod
    def get_job_description(job_id: int) -> str:
        """
        Fetches the job description text for a specific job_id.
        
        Args:
            job_id (int): The ID of the job
            
        Returns:
            str: The job description text
            
        Raises:
            HTTPException: If job_id not found or database error occurs
        """
        job_details = JobDataAccess.get_job_details(job_id)
        return job_details["job_desc"]
    
    @staticmethod
    def get_years_of_experience(job_id: int) -> int:
        """
        Fetches the years of experience required for a specific job_id.
        
        Args:
            job_id (int): The ID of the job
            
        Returns:
            int: The required years of experience
            
        Raises:
            HTTPException: If job_id not found or database error occurs
        """
        job_details = JobDataAccess.get_job_details(job_id)
        return job_details["years_of_exp"]


#  -------------------------- Access JOB_SKILLS Table From Database --------------------------


class JobSkillsDataAccess:
    def __init__(self):
        self.job_skills = None
        self.job_id = None

    def fetch_and_store_job_skills(
        self, db_conn: DatabaseConnection, job_id: int
    ) -> List[Dict]:
        """Fetches and stores skill IDs and skill names based on job_id."""
        try:
            query = """
                SELECT js.skill_id, s.skill_name
                FROM public.job_skills js
                JOIN public.skills s ON js.skill_id = s.skill_id
                WHERE js.job_id = %s
            """
            db_conn.cur.execute(query, (job_id,))

            # Fetching column names and rows
            rows = db_conn.cur.fetchall()

            if not rows:
                raise HTTPException(
                    status_code=404, detail=f"No skills found for job_id {job_id}."
                )

            # Storing the fetched data in instance variables
            self.job_id = job_id
            self.job_skills = [
                {"skill_id": row[0], "skill_name": row[1]} for row in rows
            ]
            print(f"Successfully fetched job skills for job_id {job_id}")
            print(self.job_skills)
            return self.job_skills

        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Database error while fetching job skills: {str(e)}",
            )

    def get_stored_job_skills(self) -> Optional[List[Dict]]:
        """Returns the stored job skills if available."""
        return self.job_skills

    def get_stored_job_id(self) -> Optional[int]:
        """Returns the stored job ID if available."""
        return self.job_id


#  -------------------------- Access SKILLS Table --------------------------


class SkillDataAccess:
    @staticmethod
    def get_skill_details(db_conn: DatabaseConnection) -> List[Dict]:
        """Fetches all skill details from the skills table."""
        try:
            query = """
                SELECT skill_id, skill_name
                FROM public.skills
            """
            db_conn.cur.execute(query)

            columns = [desc[0] for desc in db_conn.cur.description]
            rows = db_conn.cur.fetchall()

            if not rows:
                raise HTTPException(
                    status_code=404, detail="No skills found in the database."
                )

            skills = [dict(zip(columns, row)) for row in rows]
            return skills
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Database error while fetching skill details: {str(e)}",
            )


# -------------------------- Access JOB_SHORTLISTED_PROFILES Table --------------------------


class ShortlistDataAccess:
    @staticmethod
    def get_shortlisted_profiles(
        db_conn: DatabaseConnection, job_id: int
    ) -> List[Dict]:
        """Fetches profiles from the job_shortlisted_profiles table."""
        try:
            query = """
                SELECT shortlist_id, job_id, manager_id, candidate_id 
                FROM public.job_shortlisted_profiles
                WHERE job_id = %s
            """
            db_conn.cur.execute(query, (job_id,))

            columns = [desc[0] for desc in db_conn.cur.description]
            rows = db_conn.cur.fetchall()

            if not rows:
                raise HTTPException(
                    status_code=404,
                    detail=f"No shortlisted profiles found for job_id {job_id}. Please ensure candidates have been shortlisted for this job.",
                )

            profiles = [dict(zip(columns, row)) for row in rows]
            return profiles
        except HTTPException:
            raise
        except Exception as e:
            error_message = f"Database error while fetching shortlisted profiles: {str(e)}"
            print(error_message)  # Add direct console logging for debugging
            raise HTTPException(
                status_code=500, 
                detail=error_message
            )

# -------------------------- Access CANDIDATE Table --------------------------


class CandidateDataAccess:
    @staticmethod
    def get_candidate_details(
        db_conn: DatabaseConnection, candidate_id: int
    ) -> List[Dict]:
        """Fetches candidate details from the candidate table."""
        try:
            query = """
                SELECT candidate_id, s3_key_resume
                FROM public.candidate
                WHERE candidate_id = %s AND is_deleted = false
            """
            db_conn.cur.execute(query, (candidate_id,))

            columns = [desc[0] for desc in db_conn.cur.description]
            rows = db_conn.cur.fetchall()

            if not rows:
                raise HTTPException(
                    status_code=404, detail=f"Candidate ID {candidate_id} not found."
                )

            candidates = [dict(zip(columns, row)) for row in rows]
            return candidates
        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=500,
                detail=f"Database error while fetching candidate details: {str(e)}",
            )

# -------------------------- Access RESUME_SCORE Table --------------------------

class ScoringConfigDataAccess:
    """
    Class to access scoring configuration from the database
    """
    @staticmethod
    def get_scoring_configs(db_conn: DatabaseConnection) -> Dict[int, Dict[str, Any]]:
        """
        Fetches scoring configuration parameters from the resume_score_calculator table.
        
        Args:
            db_conn: DatabaseConnection instance
            
        Returns:
            Dict[int, Dict[str, Any]]: Dictionary mapping criteria_id to scoring parameters
        """
        try:
            query = """
                SELECT criteria_id, criteria__calc_desc, criteria__calc_formula, 
                       start_value, increment_by, max_value
                FROM public.resume_score_calculator
                ORDER BY criteria_id
            """
            db_conn.cur.execute(query)
            
            columns = [desc[0] for desc in db_conn.cur.description]
            rows = db_conn.cur.fetchall()
            
            if not rows:
                logging.warning("No scoring configuration found in resume_score_calculator table")
                return {}
            
            # Create a dictionary of scoring configs mapped by criteria_id
            scoring_configs = {}
            for row in rows:
                row_dict = dict(zip(columns, row))
                criteria_id = row_dict['criteria_id']
                scoring_configs[criteria_id] = row_dict
                
            return scoring_configs
            
        except Exception as e:
            logging.error(f"Error fetching scoring configuration: {str(e)}")
            logging.error(traceback.format_exc())
            return {}

# -------------------------- Document Parsing --------------------------


import os
import pathlib
from typing import Optional
import fitz  # PyMuPDF
from docx import Document


import os
import pathlib
from typing import Optional, List, Dict, Any

import fitz  # PyMuPDF
from docx import Document
import docx2txt
import textract
import logging

# Configure logging
logging.basicConfig(level=logging.INFO, 
                    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class DocumentParser:
    @staticmethod
    def read_file(file_path: str) -> Optional[str]:
        """
        Read file based on its extension using the appropriate method.

        Args:
            file_path (str): Path to the file to read.

        Returns:
            Optional[str]: Text content of the file or None if an error occurs.
        """
        try:
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")

            file_extension = pathlib.Path(file_path).suffix.lower()

            if file_extension == ".pdf":
                return DocumentParser.read_pdf(file_path)
            elif file_extension in [".docx", ".doc"]:
                return DocumentParser.read_docx(file_path)
            elif file_extension == ".txt":
                return DocumentParser.read_txt(file_path)
            else:
                # Try with textract for other formats
                try:
                    return DocumentParser.read_with_textract(file_path)
                except:
                    raise ValueError(f"Unsupported file format: {file_extension}")
        except Exception as e:
            logger.error(f"Error reading file: {str(e)}")
            return None

    @staticmethod
    def read_pdf(file_path: str) -> Optional[str]:
        """
        Extract text from a PDF file using PyMuPDF (fitz) with enhanced extraction.
        
        This method ensures extraction of all text elements including:
        - Regular text blocks
        - Text in tables
        - Headers and footers
        - Annotations and form fields
        """
        try:
            doc = fitz.open(file_path)
            full_text = []
            
            for page_num, page in enumerate(doc):
                # Extract text with more inclusive parameters
                text = page.get_text("text")
                
                # Add page delimiter for better structure
                page_text = f"\n--- Page {page_num + 1} ---\n{text}"
                
                # Extract annotations (comments, highlights, etc.)
                annots = page.annots()
                if annots:
                    annot_text = "\nAnnotations:\n"
                    for annot in annots:
                        if annot.info.get("content"):
                            annot_text += f"- {annot.info['content']}\n"
                    page_text += annot_text
                
                # Try to get text from form fields if any
                widgets = page.widgets()
                if widgets:
                    form_text = "\nForm Fields:\n"
                    for widget in widgets:
                        if widget.text:
                            form_text += f"- {widget.field_name}: {widget.text}\n"
                    page_text += form_text
                
                full_text.append(page_text)
            
            doc.close()  # Properly close the document
            
            # Fall back to alternative extraction if minimal text was found
            combined_text = "\n".join(full_text).strip()
            if len(combined_text) < 100:  # Arbitrary threshold
                logger.warning("Minimal text extracted with primary method, trying fallback...")
                try:
                    # Try with textract as fallback
                    fallback_text = DocumentParser.read_with_textract(file_path)
                    if fallback_text and len(fallback_text) > len(combined_text):
                        return fallback_text
                except Exception as e:
                    logger.warning(f"Fallback extraction failed: {str(e)}")
            
            return combined_text
        except Exception as e:
            logger.error(f"Error reading PDF file: {str(e)}")
            # Try fallback method
            try:
                return DocumentParser.read_with_textract(file_path)
            except:
                return None

    @staticmethod
    def read_docx(file_path: str) -> Optional[str]:
        """
        Extract text from a DOCX/DOC file with enhanced extraction.
        
        This method ensures extraction of all text elements including:
        - Paragraphs
        - Tables
        - Headers and footers
        - Text boxes and other shapes
        """
        try:
            # Try docx2txt first (better at extracting tables and other elements)
            text = docx2txt.process(file_path)
            
            # If docx2txt didn't extract much, try python-docx
            if not text or len(text.strip()) < 100:  # Arbitrary threshold
                doc = Document(file_path)
                
                # Get text from paragraphs
                paragraphs_text = [paragraph.text for paragraph in doc.paragraphs]
                
                # Extract text from tables
                tables_text = []
                for table in doc.tables:
                    for row in table.rows:
                        row_text = [cell.text for cell in row.cells]
                        tables_text.append(" | ".join(row_text))
                
                # Extract text from headers and footers if available
                headers_footers = []
                try:
                    for section in doc.sections:
                        # Header
                        header = section.header
                        if header:
                            for paragraph in header.paragraphs:
                                if paragraph.text.strip():
                                    headers_footers.append(f"Header: {paragraph.text}")
                        
                        # Footer
                        footer = section.footer
                        if footer:
                            for paragraph in footer.paragraphs:
                                if paragraph.text.strip():
                                    headers_footers.append(f"Footer: {paragraph.text}")
                except:
                    pass  # Headers/footers extraction is optional
                
                # Combine all text elements
                combined_text = (
                    "\n\n".join(paragraphs_text) + 
                    ("\n\n=== TABLES ===\n" + "\n".join(tables_text) if tables_text else "") +
                    ("\n\n=== HEADERS & FOOTERS ===\n" + "\n".join(headers_footers) if headers_footers else "")
                )
                
                text = combined_text
            
            # If we still don't have much text, try textract as a last resort
            if not text or len(text.strip()) < 100:
                try:
                    textract_text = DocumentParser.read_with_textract(file_path)
                    if textract_text and len(textract_text) > len(text):
                        return textract_text
                except:
                    pass
                
            return text.strip()
        except Exception as e:
            logger.error(f"Error reading DOCX file: {str(e)}")
            # Try fallback method
            try:
                return DocumentParser.read_with_textract(file_path)
            except:
                return None

    @staticmethod
    def read_txt(file_path: str) -> Optional[str]:
        """Read text from a TXT file with multiple encoding support."""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                with open(file_path, "r", encoding=encoding) as file:
                    return file.read().strip()
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.error(f"Error reading TXT file with {encoding}: {str(e)}")
        
        # If all encodings fail, try binary mode
        try:
            with open(file_path, 'rb') as file:
                binary_content = file.read()
                # Try to decode with error handling
                return binary_content.decode('utf-8', errors='replace').strip()
        except Exception as e:
            logger.error(f"Error reading TXT file in binary mode: {str(e)}")
            return None

    @staticmethod
    def read_with_textract(file_path: str) -> Optional[str]:
        """Use textract as a fallback for unsupported file types."""
        try:
            text = textract.process(file_path).decode('utf-8')
            return text.strip()
        except Exception as e:
            logger.error(f"Error extracting text with textract: {str(e)}")
            return None

    def parse_resume(self, file_path: str) -> Optional[str]:
        """
        Parse a resume file and return the complete text content.
        
        Args:
            file_path (str): Path to the resume file.
            
        Returns:
            Optional[str]: The complete text content of the resume or None if an error occurs.
        """
        text = self.read_file(file_path)
        
        # If we got text back, verify it has substantial content
        if text and len(text.strip()) < 50:  # Very short text might indicate extraction issues
            logger.warning(f"Extracted content seems unusually short ({len(text)} chars). Trying alternate methods.")
            try:
                # Try textract as a last resort
                alternate_text = self.read_with_textract(file_path)
                if alternate_text and len(alternate_text) > len(text):
                    return alternate_text
            except:
                pass
        
        return text


# -------------------------- Skill Matcher With Database -------------------------


class SkillMatcher:
    def __init__(self, job_id: int, api_key: str):
        self.job_id = job_id
        self.api_key = api_key
        self.skills_fetcher = SkillsFetcher()
        self.resume_analyzer = ResumeAnalyzer(api_key)

    def match_skills_and_get_score(self, resume_text: str, job_description_text: str):
        # Step 1: Fetch the required skills from the SkillsFetcher
        self.skills_fetcher.fetch_skills(self.job_id)
        required_skills = (
            self.skills_fetcher.get_skills_dict()
        )  # Dictionary of skill_id: skill_name

        # Step 2: Analyze the resume using ResumeAnalyzer
        skill_analysis_results = self.resume_analyzer.analyze_resume_skills(
            resume_text, required_skills
        )

        if isinstance(skill_analysis_results, str):  # If it's an error message
            return skill_analysis_results

        # Step 3: Extract skills from resume analysis
        resume_skills = []
        if "skills_analysis" in skill_analysis_results:
            for skill_data in skill_analysis_results["skills_analysis"]:
                resume_skills.append(skill_data["skill"])

        # Step 4: Match skills from both required skills and analyzed skills
        matched_skills = {}
        for skill_id, skill_name in required_skills.items():
            if skill_name in resume_skills:
                # Step 5: Calculate the score for the matched skills using ResumeAnalyzer's method
                mentions = self.resume_analyzer.count_skill_occurrences(
                    resume_text, skill_name
                )
                last_used_year = max(
                    [
                        year
                        for start, end, _ in self.resume_analyzer.extract_skill_duration_from_experience(
                            resume_text, skill_name
                        )
                        for year in [start, end]
                    ],
                    default=self.resume_analyzer.current_year,
                )
                total_experience = self.resume_analyzer.calculate_total_skill_duration(
                    self.resume_analyzer.extract_skill_duration_from_experience(
                        resume_text, skill_name
                    )
                )

                scores = self.resume_analyzer.calculate_skill_scores(
                    mentions, last_used_year, total_experience
                )
                matched_skills[skill_name] = {
                    "mentions": mentions,
                    "last_used_year": last_used_year,
                    "total_experience": total_experience,
                    "frequency_score": scores["frequency_score"],
                    "recency_score": scores["recency_score"],
                    "duration_score": scores["duration_score"],
                    "total_score": scores["total_score"],
                }

        return matched_skills


# -------------------------- Required Skills Feteching From Database --------------------------


class SkillsFetcher:
    """
    A simplified class to fetch and store job skills and their IDs.
    """

    def __init__(self):
        """
        Initialize the fetcher with variables to store skills
        """
        self.skill_ids = []  # List to store skill IDs
        self.skill_names = []  # List to store skill names
        self.skills_dict = {}  # Dictionary to store skill_id: skill_name pairs

    def fetch_skills(self, job_id: int) -> None:
        """
        Fetch skills and their IDs for a specific job ID and store them in instance variables.

        Args:
            job_id (int): The ID of the job to fetch skills for
        """
        with DatabaseConnection() as db_conn:
            try:
                query = """
                    SELECT 
                        js.skill_id,
                        s.skill_name
                    FROM public.job_skills js
                    JOIN public.skills s ON js.skill_id = s.skill_id
                    WHERE js.job_id = %s AND js.is_deleted = false
                """

                db_conn.cur.execute(query, (job_id,))
                results = db_conn.cur.fetchall()

                if not results:
                    raise Exception(f"No skills found for job ID {job_id}")

                # Clear existing data
                self.skill_ids.clear()
                self.skill_names.clear()
                self.skills_dict.clear()

                # Store results in instance variables
                for row in results:
                    skill_id, skill_name = row
                    self.skill_ids.append(skill_id)
                    self.skill_names.append(skill_name)
                    self.skills_dict[skill_id] = skill_name

            except Exception as e:
                print(f"Error fetching skills: {str(e)}")
                raise

    def get_skill_ids(self) -> list:
        """Return the list of skill IDs."""
        return self.skill_ids

    def get_skill_names(self) -> list:
        """Return the list of skill names."""
        return self.skill_names

    def get_skills_dict(self) -> dict:
        """Return the dictionary of skill_id: skill_name pairs."""
        return self.skills_dict


# -------------------------- Resume Analyzer --------------------------

import anthropic
import anthropic
import logging
import re
import traceback
from datetime import datetime
from typing import List, Optional, Dict, Set, Any, Union, Tuple

class ResumeAnalyzer:
    def __init__(self, api_key: str):
        """Initialize the ResumeAnalyzer with Anthropic API key"""
        if not api_key:
            raise ValueError("API key cannot be empty")
        
        # Initialize with explicit parameters, avoiding proxy-related parameters
        self.client = anthropic.Anthropic(
            api_key=api_key,
            base_url="https://api.anthropic.com",
        )
        
        self.current_year = datetime.now().year
        self.model = "claude-3-7-sonnet-20250219"  # Use a stable model name without date suffix
        self.scoring_configs = None  # Will be loaded when needed
        self.skills_fetcher = None  # Will be set if needed for skills-related criteria

    def load_scoring_configs(self, db_conn) -> None:
        """Load scoring configurations from the database"""
        self.scoring_configs = ScoringConfigDataAccess.get_scoring_configs(db_conn)
        logging.info(f"Loaded {len(self.scoring_configs)} scoring configurations from database")

    def get_skill_scoring_params(self, criteria_id: int) -> Tuple[float, float, float]:
        """
        Get scoring parameters for a specific criteria from the loaded config.
        
        Returns:
            Tuple[float, float, float]: (start_value, increment_by, max_value)
        """
        if not self.scoring_configs:
            # Default values if no configs are loaded (5 points per occurrence, max 100)
            return 0, 5, 100
            
        config = self.scoring_configs.get(criteria_id)
        if not config:
            # Default values if specific criteria config not found
            return 0, 5, 100
            
        return (
            float(config.get('start_value', 0)),
            float(config.get('increment_by', 5)), 
            float(config.get('max_value', 100))
        )

    def calculate_skill_score(self, occurrences: int, criteria_id: int) -> int:
        """
        Calculate skill score based on occurrences and database configuration.
        
        Args:
            occurrences: Number of times the skill appears in the resume
            criteria_id: ID of the scoring criteria to use for calculation
            
        Returns:
            int: Calculated score capped at max_value
        """
        start_value, increment_by, max_value = self.get_skill_scoring_params(criteria_id)
        score = start_value + (occurrences * increment_by)
        return min(int(score), int(max_value))

    def analyze_resume_skills(self, resume_text: str, required_skills: Dict[str, str]) -> Union[Dict[int, Dict[str, Any]], str]:
        """Analyze skills in resume using enhanced regex matching and AI validation"""
        import re
        
        if not resume_text or not required_skills:
            return "Error: Missing resume text or required skills"

        # Map skill indexes to criteria IDs
        skill_criteria_mapping = {
            0: 1,  # skill1 -> criteria_id 1
            1: 2,  # skill2 -> criteria_id 2
            2: 3,  # skill3 -> criteria_id 3
            3: 9,  # skill4 -> criteria_id 9
            4: 8,  # skill5 -> criteria_id 8
        }

        # Initialize criteria scores
        criteria_scores = {
            criteria_id: {"criteria_id": criteria_id, "name": f"skill{skill_idx + 1}", "score": 0}
            for skill_idx, criteria_id in skill_criteria_mapping.items()
        }
        
        # Create regex patterns for each skill with improved handling
        skills_list = list(required_skills.values())
        skill_patterns = []
        
        # First, build a dictionary to identify potential substring conflicts
        # E.g., if skills include both "Java" and "JavaScript", we need special handling
        skill_conflicts = {}
        for i, skill_a in enumerate(skills_list):
            for j, skill_b in enumerate(skills_list):
                if i != j and skill_a.lower() in skill_b.lower():
                    if skill_a.lower() not in skill_conflicts:
                        skill_conflicts[skill_a.lower()] = []
                    skill_conflicts[skill_a.lower()].append(skill_b.lower())
        
        for skill in skills_list:
            # Handle multi-word skills
            words = skill.split()
            
            # Check if this skill might be matched as part of another skill
            if skill.lower() in skill_conflicts:
                # Create a pattern that won't match when this skill is part of another skill
                # First, create the base skill pattern
                base_pattern = r'\b' + re.escape(skill) + r'\b'
                
                # Then, add negative lookaheads for each potential conflict
                # This ensures "Java" won't match in "JavaScript"
                negative_lookaheads = []
                for conflict in skill_conflicts[skill.lower()]:
                    # Get the part that follows after the current skill
                    suffix = conflict[len(skill.lower()):]
                    if suffix:
                        # The pattern should match the skill not followed by the suffix
                        # Case insensitive by using character classes for each letter
                        suffix_pattern = ''.join(f'[{c.lower()}{c.upper()}]' for c in suffix)
                        negative_lookaheads.append(f'(?!{suffix_pattern})')
                
                # Combine the base pattern with all negative lookaheads
                pattern = base_pattern + ''.join(negative_lookaheads)
            elif len(words) > 1:
                # Pattern for exact match with possible punctuation/space variations
                exact_pattern = r'\b' + r'[\s\-_]*'.join(re.escape(word) for word in words) + r'\b'
                # Pattern for common abbreviations (first letters)
                abbrev = ''.join(word[0] for word in words)
                abbrev_pattern = r'\b' + re.escape(abbrev) + r'\b'
                # Pattern for joined words (no spaces)
                joined_pattern = r'\b' + re.escape(''.join(words)) + r'\b'
                # Combined pattern
                pattern = f"({exact_pattern}|{abbrev_pattern}|{joined_pattern})"
            else:
                # Single word pattern
                pattern = r'\b' + re.escape(skill) + r'\b'
            
            # Add pattern matching for text in parentheses
            parentheses_pattern = r'\(' + re.escape(skill) + r'\)'
            bracketed_pattern = r'\[' + re.escape(skill) + r'\]'
            
            # Combined pattern including parentheses/brackets options
            combined_pattern = f"({pattern}|{parentheses_pattern}|{bracketed_pattern})"
            
            skill_patterns.append(combined_pattern)
        
        # Create skills_text for the prompt
        skills_text = "\n".join([f"{idx + 1}. {skill}" for idx, skill in enumerate(skills_list)])
        
        # Pre-count with regex (case insensitive)
        regex_counts = []
        for pattern in skill_patterns:
            count = len(re.findall(pattern, resume_text, re.IGNORECASE))
            regex_counts.append(count)
        
        # Identify potential skill conflicts (where one skill name is part of another)
        conflict_pairs = []
        for i, skill_a in enumerate(skills_list):
            for j, skill_b in enumerate(skills_list):
                if i != j and skill_a.lower() in skill_b.lower():
                    conflict_pairs.append(f"{skill_a} vs {skill_b}")
        
        # Create conflict warning text if needed
        conflict_warning = ""
        if conflict_pairs:
            conflict_warning = "IMPORTANT - These skills need careful differentiation:\n"
            conflict_warning += "\n".join([f"- {pair}" for pair in conflict_pairs])
        
        # Create prompt with regex counts as reference and specific instructions for tables and parentheses
        prompt = f"""
        Review this resume for specific technical skills. I've already performed basic regex matching and found preliminary counts, but need you to validate and refine these results.

        Skills to analyze and preliminary regex counts:
        {', '.join([f"{skills_list[i]} (regex found: {regex_counts[i]})" for i in range(len(skills_list))])}

        For each skill, consider variations including:
        - Different capitalizations
        - Space/hyphen variations 
        - Common abbreviations
        - Related components/frameworks
        - Domain-specific terminology
        - Check in tables listed in resume
        - Check skills mentioned in parentheses like (AWS) or [AWS]
        
        IMPORTANT:
        1. Pay special attention to skills listed in tables or bullets
        2. Count skills mentioned within parentheses or brackets like (AWS) or [Python]
        3. Look for skills in section headers, bulleted lists, and table cells
        
        {conflict_warning}
        
        CRITICAL: For all skills, ensure you count exact matches only. For example, "Java" should not be counted when it appears as part of "JavaScript" and vice versa.

        Resume text:
        {resume_text}  # Trim if needed

        For each skill (numbered 1-{len(skills_list)}), provide ONLY:
        - The exact count of occurrences
        - A score calculated as (occurrences × increment) + start_value, capped at max_value
        - Note: Different skills might have different scoring parameters

        Format each response as:
        Skill 1: Found: [count] | Score: [calculated score]
        Skill 2: Found: [count] | Score: [calculated score]
        ...
        """

        try:
            # Call Claude API
            response = self.client.messages.create(
                model=self.model,
                max_tokens=1000,
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
            )

            # Parse response text
            raw_text = response.content[0].text
            
            # Use regex to extract counts and scores
            result_pattern = r'Skill\s+(\d+):\s+Found:\s+(\d+)\s*\|\s*Score:\s+(\d+)'
            matches = re.findall(result_pattern, raw_text)
            
            raw_scores = {}
            for match in matches:
                skill_num, count_str, score_str = match
                skill_key = f"Skill {skill_num}"
                try:
                    count = int(count_str)
                    # Instead of trusting the AI's score calculation, we'll recalculate based on our parameters
                    skill_idx = int(skill_num) - 1
                    criteria_id = skill_criteria_mapping.get(skill_idx)
                    score = self.calculate_skill_score(count, criteria_id) if criteria_id else 0
                    
                    raw_scores[skill_key] = {
                        'occurrences': count,
                        'score': score
                    }
                except (ValueError, TypeError) as e:
                    logging.warning(f"Error parsing skill result: {e}")
            
            # If AI results are missing, use regex counts as fallback
            for i in range(len(skills_list)):
                skill_key = f"Skill {i+1}"
                if skill_key not in raw_scores:
                    count = regex_counts[i]
                    criteria_id = skill_criteria_mapping.get(i)
                    score = self.calculate_skill_score(count, criteria_id) if criteria_id else 0
                    
                    raw_scores[skill_key] = {
                        'occurrences': count,
                        'score': score
                    }
            
            # Update scores for skills
            for skill_idx, skill_name in enumerate(skills_list):
                if skill_idx in skill_criteria_mapping and skill_idx < len(skills_list):
                    criteria_id = skill_criteria_mapping[skill_idx]
                    skill_data = raw_scores.get(f"Skill {skill_idx + 1}", {'occurrences': 0, 'score': 0})
                    criteria_scores[criteria_id]["score"] = skill_data['score']
                    # Add occurrence data for reference
                    criteria_scores[criteria_id]["occurrences"] = skill_data['occurrences']
            
            return criteria_scores
            
        except Exception as e:
            logging.error(f"Error in skills analysis: {str(e)}")
            logging.error(traceback.format_exc())
            return f"Error in skills analysis: {str(e)}"

    def _job_description_requires_education(self, job_description: str) -> bool:
        """
        Check if job description mentions education requirements
        
        Args:
            job_description (str): The job description text
            
        Returns:
            bool: True if education appears to be required, False otherwise
        """
        education_keywords = [
            r"degree", r"bachelor", r"master", r"phd", r"education", 
            r"university", r"college", r"diploma", r"certification"
        ]
        
        for keyword in education_keywords:
            if re.search(rf"\b{keyword}s?\b", job_description, re.IGNORECASE):
                return True
        
        return False
            
    def general_analyze_resume(
            self, resume_text: str, job_description: str, job_id: int, job_years_exp: int = None
        ) -> Union[Dict[str, Any], str]:
            """
            Analyze resume based on non-skill criteria using database-driven scoring configuration.

            Args:
                resume_text (str): The text content of the resume
                job_description (str): The job description text
                job_id (int): The job ID
                job_years_exp (int, optional): Years of experience required in the job

            Returns:
                Union[Dict[str, Any], str]: Dictionary containing analysis results or error message
            """
            # Log the start of general resume analysis
            logging.info(f"Starting general resume analysis for job ID {job_id}")
            
            # Create a more dynamic prompt based on database configuration
            criteria_prompts = []
            
            # List of criteria IDs that we'll be analyzing
            criteria_ids = [4, 5, 6, 7, 10, 11, 12, 13, 14, 15]
            
            # Log the criteria IDs being analyzed
            logging.info(f"Analyzing resume against {len(criteria_ids)} criteria: {criteria_ids}")
            
            # Get scoring parameters for each criteria
            for criteria_id in criteria_ids:
                start_value, increment_by, max_value = self.get_skill_scoring_params(criteria_id)
                logging.info(f"Criteria {criteria_id} scoring parameters: start={start_value}, increment={increment_by}, max={max_value}")
                
                if criteria_id == 4:  # Education Match with Job Description
                    criteria_prompts.append(f"""
                    4. Education Match:
                    Critical thing to consider first if degree is required or preferred in {job_description} then check if education or degree credentials are present in {resume_text} 
                        - If NO education or degree is found in the resume:
                            - If education or degree IS required in {job_description}, assign {start_value}
                            - If education or degree is NOT required in {job_description}, assign {max_value}
                        - If education or degree IS found in the resume:
                            - If it matches the requirements in {job_description}, assign {max_value}
                            - If it doesn't match the requirements, assign {start_value}
                    """)

                
                elif criteria_id == 5:  # Job Title Match
                    criteria_prompts.append(f"""
                    5. Job Title Match:
                    - First, identify the EXACT job title from the job description:For example "Java Lead Developer"
                    - Next, identify the candidate's most recent/current job title from their resume:For example "Senior Java Full Stack developer or Java Tech Lead "
                    - Use this SPECIFIC hierarchy for evaluation (from lowest to highest):
                    Intern < Junior/Associate < (no prefix) < Senior < Lead < Manager < Director < VP < C-level
                    - Score EXACTLY {max_value} ONLY if:
                    * The candidate's last job title is IDENTICAL to the job description title example "Java Lead Developer" or "Java Tech Lead"
                    * The candidate's last job title is at a HIGHER level in the hierarchy than the job description title
                    - Score EXACTLY {start_value} if:
                    * The candidate's last job title is at a LOWER level in the hierarchy than the job description title
                    * The candidate's last job title is in a different role category entirely
                    - IMPORTANT: In this case, "Lead" is HIGHER in the hierarchy than "Senior", so a "Senior" role should receive {start_value} when applying for a "Lead" role
                    - Focus ONLY on the candidate's CURRENT/MOST RECENT position title
                    """)
                
                elif criteria_id == 6:  # Total IT Experience
                    criteria_prompts.append(f"""
                    6. Total IT Experience:
                    - Required number of years of IT Experience: {job_years_exp}
                    - Calculate the total number of years the candidate has worked get total years in IT or professional experience as mentioned in the resume.
                    - For every additional year beyond the baseline required {job_years_exp}, give {increment_by} points.
                    - Cap the points at {max_value}.
                    
                    Examples:
                    * {job_years_exp} years required, {job_years_exp + 2} years in resume = {min(2 * increment_by, max_value)} points
                    * {job_years_exp} years required, {job_years_exp + 4} years in resume = {min(4 * increment_by, max_value)} points
                    * Many additional years = {max_value} points (capped)
                    """)
                
                elif criteria_id == 7:  # Current Skills Usage
                    # Get the first skill from the skills fetcher
                    # We need to make sure we've already fetched skills for this job
                    if not hasattr(self, 'skills_fetcher') or not self.skills_fetcher:
                        # Create a skills fetcher if we don't have one
                        self.skills_fetcher = SkillsFetcher()
                        # Fetch skills for the current job
                        self.skills_fetcher.fetch_skills(job_id)
                    
                    # Get the list of skill names
                    skill_names = self.skills_fetcher.get_skill_names()
                    # Get the first skill if available
                    first_skill = skill_names[0] if skill_names else "required skill"
                    logging.info(f"Current Skills Usage - using first skill: {first_skill}")
                    
                    criteria_prompts.append(f"""
                    7. Current Skills Usage:
                    - Focus ONLY on this specific skill: "{first_skill}"
                    - Score {max_value} if "{first_skill}" is used in candidate's current/last project or last job role
                    - Score {start_value} if "{first_skill}" is not used in current/last project or last job role
                    - Base your evaluation on explicit mentions of this skill in relation to recent work
                    - If the skill is not clearly mentioned in connection with recent work, score {start_value}
                    """)
                
                elif criteria_id == 10:  # Domain Experience
                    criteria_prompts.append(f"""
                    10. Domain Experience:
                    - Identify the specific domain required in the {job_description} (e.g., banking, healthcare, retail, manufacturing, etc.)
                    - Review each company/employer listed in the candidate's resume and determine if they belong to the same domain required in the job description
                    - For companies that match the required domain, calculate the candidate's total years of experience
                    - Add EXACTLY {increment_by} points for EACH FULL YEAR of experience in companies from the matching domain
                    - Examples:
                    * If JD requires banking domain experience, and candidate worked at Bank of America for 3 years and Chase Bank for 2 years = 5 years domain experience = {min(5 * increment_by, max_value)} points
                    * If JD requires healthcare domain, but candidate only worked in retail = 0 points
                    * If candidate has 10+ years in matching domain = {max_value} points (capped at max {max_value})
                    - Score {max_value} if no specific domain is mentioned in the job description
                    - Consider company descriptions, project details, or client information when determining domain relevance
                    - Do not award points for companies/employers outside the required domain
                    - Do not award partial points for partial years
                    """)
                
                elif criteria_id == 11:  # Awards and Recognition
                    criteria_prompts.append(f"""
                    11. Awards and Recognition:
                    - Score {max_value} if any awards present
                    - Score {start_value} if none
                    """)
                
                elif criteria_id == 12:  # Certifications
                    criteria_prompts.append(f"""
                    12. Certifications:
                    - Score {max_value} if any certification present
                    - Score {start_value} if none
                    """)
                
                elif criteria_id == 13:  # Education
                    criteria_prompts.append(f"""
                    13. Education:
                    -                        
                    - Score {max_value} for CS/CE degree 
                    - Score {start_value} if IT education is not mentioned in resume
                    - Score {start_value} for others
                    """)
                
                elif criteria_id == 14:  # Skill Usage Duration
                    # Get the first skill from the skills fetcher
                    # We need to make sure we've already fetched skills for this job
                    if not hasattr(self, 'skills_fetcher') or not self.skills_fetcher:
                        # Create a skills fetcher if we don't have one
                        self.skills_fetcher = SkillsFetcher()
                        # Fetch skills for the current job
                        self.skills_fetcher.fetch_skills(job_id)
                    
                    # Get the list of skill names
                    skill_names = self.skills_fetcher.get_skill_names()
                    # Get the first skill if available
                    first_skill = skill_names[0] if skill_names else "required skill"
                    logging.info(f"Skill Usage Duration - using first skill: {first_skill}")
                    
                    criteria_prompts.append(f"""
                    14. Skill Usage Duration:
                    - Focus ONLY on this specific primary skill: "{first_skill}"
                    - Award EXACTLY {increment_by} points per year the candidate has used "{first_skill}"
                    - Calculate based on total relevant work experiences or job roles with "{first_skill}", not exceeding actual years of experience
                    - Examples:
                    * 1 year using "{first_skill}" = {increment_by} points
                    * 10 years using "{first_skill}" = {min(10 * increment_by, max_value)} points
                    * 20+ years using "{first_skill}" = {max_value} points (capped at {max_value})
                    - Be precise in counting only experience with this specific skill, not general experience
                    """)
                
                elif criteria_id == 15:  # Average Project Length
                    criteria_prompts.append(f"""
                    15. Average Project Length:
                    - Calculate the average number of years per project OR per employer from the resume
                    - For each position or project in the work experience section, calculate its duration in years
                    - Sum all durations and divide by the total number of projects/employers to get the average
                    - If incomplete dates are provided, make a reasonable estimate based on context
                    - Score EXACTLY {max_value} if the average is 3 or more years
                    - Score EXACTLY {start_value} if the average is less than 3 years
                    - Provide the exact average in your explanation (e.g., "Average is 2.7 years")
                    """)
            
            # Combine all criteria prompts
            all_criteria_prompts = "\n".join(criteria_prompts)
            
            # Create the full prompt
            prompt = f"""
            Analyze the following resume against the job description and provide scores for each criterion.
            Provide ONLY numerical scores (0-{max_value}) for each criterion.

            Scoring Criteria:
            {all_criteria_prompts}

            Resume:
            {resume_text}  # Limit text size if needed

            Job Description:
            {job_description}  # Limit text size if needed

            Format your response exactly like this:
            4: [score]
            5: [score]
            6: [score]
            7: [score]
            10: [score]
            11: [score]
            12: [score]
            13: [score]
            14: [score]
            15: [score]
            """

            try:
                # Log the API call attempt
                logging.info("Calling Claude API for general resume analysis")
                
                # Call Claude API
                response = self.client.messages.create(
                    model=self.model,
                    max_tokens=4000,
                    messages=[{"role": "user", "content": prompt}],
                    temperature=0.3,
                )

                # Log successful API call
                logging.info("Received response from Claude API")
                
                # Parse the response using our updated _parse_general_analysis method
                criteria_scores = self._parse_general_analysis(response.content[0].text)
                
                # Post-processing fix for education score when no education is found
                if 4 in criteria_scores and criteria_scores[4].get("education_found") is False:
                    # Check if job description requires education
                    if self._job_description_requires_education(job_description):
                        criteria_scores[4]["score"] = 0  # Force score to 0 if education required but not found
                        logging.info("Forced education score to 0 because no education found but JD requires it")
                
                # Log the scores
                for criteria_id, score_data in criteria_scores.items():
                    logging.info(f"Criteria {criteria_id} ({score_data.get('name', 'unknown')}): Score {score_data.get('score', 0)}")

                # Calculate overall score using database-configured weights if available
                weighted_scores = []
                total_weight = 0
                
                # Log weights being used
                logging.info("Calculating weighted scores:")
                
                for criteria_id, score_data in criteria_scores.items():
                    # Check if there's a weight configured in the database
                    config = self.scoring_configs.get(criteria_id, {}) if self.scoring_configs else {}
                    weight = float(config.get('weight', 1.0)) if 'weight' in config else 1.0
                    
                    # Log the weight being used
                    logging.info(f"  Criteria {criteria_id}: score={score_data['score']}, weight={weight}")
                    
                    weighted_scores.append(score_data["score"] * weight)
                    total_weight += weight
                
                # Calculate the weighted average or simple average if no weights
                if total_weight > 0:
                    overall_score = sum(weighted_scores) / total_weight
                    logging.info(f"Calculated weighted average score: {overall_score:.2f}")
                else:
                    # Fallback to simple average
                    scores = [score_data["score"] for score_data in criteria_scores.values()]
                    overall_score = sum(scores) / len(scores) if scores else 0
                    logging.info(f"Calculated simple average score (fallback): {overall_score:.2f}")
                
                logging.info("General resume analysis completed successfully")
                return {
                    "overall_score": overall_score, 
                    "detailed_scores": criteria_scores
                }

            except Exception as e:
                logging.error(f"Error in general resume analysis: {str(e)}")
                logging.error(traceback.format_exc())
                return f"Error in analysis: {str(e)}"

    def _parse_skills_response(self, response_text: str) -> Dict[str, Dict[str, int]]:
        """Parse the skills analysis response with improved pattern matching"""
        logging.info("Parsing skills response")
        skills_scores = {}
        current_skill = None
        skill_pattern = re.compile(r"skill\s+(\d+)[:\s-]+([^:\n]+)", re.IGNORECASE)
        found_pattern = re.compile(r"found\s*:?\s*(\d+)", re.IGNORECASE)
        score_pattern = re.compile(r"(?:final\s+)?score\s*:?\s*(\d+)", re.IGNORECASE)
        
        # First try to extract all skills with their details in one pass
        skill_blocks = re.split(r"\n\s*\n|\n(?=Skill\s+\d+)", response_text, flags=re.IGNORECASE)
        logging.info(f"Found {len(skill_blocks)} skill blocks in response")
        
        for block in skill_blocks:
            block = block.strip()
            if not block:
                continue
                
            # Try to identify the skill number and name
            skill_match = skill_pattern.search(block)
            if not skill_match:
                continue
                
            skill_num = skill_match.group(1)
            skill_key = f"Skill {skill_num}"
            skills_scores[skill_key] = {"occurrences": 0, "score": 0}
            
            # Look for occurrences count
            found_match = found_pattern.search(block)
            if found_match:
                try:
                    occurrences = int(found_match.group(1))
                    skills_scores[skill_key]["occurrences"] = occurrences
                    logging.info(f"Parsed {skill_key}: Found {occurrences} occurrences")
                except ValueError:
                    logging.warning(f"Could not parse occurrence count for {skill_key}")
                    pass
            
            # Look for explicit score, but we won't use it directly
            score_match = score_pattern.search(block)
            if score_match:
                try:
                    # We're reading this but not using it - will override below
                    score = int(score_match.group(1))
                    logging.info(f"Parsed {skill_key}: LLM calculated score of {score}")
                except ValueError:
                    pass
        
        # If the block approach failed, fall back to line-by-line parsing
        if not skills_scores:
            logging.info("Block parsing failed, falling back to line-by-line parsing")
            lines = response_text.split("\n")
            for line in lines:
                line = line.strip()
                if not line:
                    continue

                skill_match = re.search(r"skill\s+(\d+)", line, re.IGNORECASE)
                if skill_match:
                    current_skill = f"Skill {skill_match.group(1)}"
                    skills_scores[current_skill] = {"occurrences": 0, "score": 0}
                    continue

                if current_skill:
                    found_match = re.search(r"found\s*:?\s*(\d+)", line, re.IGNORECASE)
                    if found_match:
                        try:
                            occurrences = int(found_match.group(1))
                            skills_scores[current_skill]["occurrences"] = occurrences
                            logging.info(f"Line parsing: {current_skill} found {occurrences} occurrences")
                        except ValueError:
                            pass
                        continue

        # Calculate the scores based on our own formula, regardless of what the LLM returned
        for skill_key, skill_data in skills_scores.items():
            occurrences = skill_data.get('occurrences', 0)
            skill_data['score'] = min(occurrences * 5, 100)
            logging.info(f"{skill_key} - Final calculation: {occurrences} occurrences -> score: {skill_data['score']}")
        
        # Add debug logging
        logging.debug(f"Parsed skills scores: {skills_scores}")
        
        return skills_scores

    def _parse_general_analysis(self, response_text: str) -> Dict[int, Dict[str, Any]]:
        """
        Parse the general analysis text into scores with detailed explanations
        
        Args:
            response_text (str): The response text from the API
            
        Returns:
            Dict[int, Dict[str, Any]]: Dictionary containing parsed criteria scores with explanations
        """
        # Log that we're starting to parse the general analysis
        logging.info("Starting to parse general analysis response")
        
        criteria_mapping = {
            4: "education_match",      # Education Match
            5: "job_title_match",      # Job Title Match (last job title)
            6: "it_experience",        # Total IT Experience
            7: "current_skills",       # Current Skills Usage (first skill)
            10: "domain_experience",   # Domain Experience (company matching)
            11: "awards",              # Awards and Recognition
            12: "certifications",      # Certifications
            13: "education",           # Education Type (CS/CE)
            14: "skill_duration",      # Skill Usage Duration (first skill)
            15: "project_length",      # Average Project Length
        }
        
        logging.info(f"Criteria mapping: {criteria_mapping}")
        
        # Initialize results with all expected criteria
        results = {
            criteria_id: {
                "criteria_id": criteria_id,
                "name": name, 
                "score": 0,
                "explanation": "",
                "education_found": True  # Default to true, will be set to False if no education found
            } 
            for criteria_id, name in criteria_mapping.items()
        }
        
        logging.info(f"Initialized results structure with {len(results)} criteria")
        
        # First try to identify score blocks for each criterion
        score_blocks = {}
        current_block = None
        current_text = []
        
        logging.info("Starting block-based parsing approach")
        
        # Split by lines and look for criterion headers
        lines = response_text.split("\n")
        logging.info(f"Response split into {len(lines)} lines")
        
        for i, line in enumerate(lines):
            line = line.strip()
            if not line:
                continue
                
            # Check if this line starts a new criterion block
            criterion_header_match = re.search(r"^\s*(?:criterion\s*)?(\d+)[:.]\s*(\d+)", line, re.IGNORECASE)
            if criterion_header_match:
                # If we were building a previous block, save it
                if current_block is not None and current_text:
                    score_blocks[current_block] = "\n".join(current_text)
                    logging.info(f"Saved block for criteria {current_block}, length: {len(current_text)} lines")
                    
                # Start a new block
                current_block = int(criterion_header_match.group(1))
                score = int(criterion_header_match.group(2))
                
                # Update the score in our results
                if current_block in criteria_mapping:
                    results[current_block]["score"] = min(score, 100)
                    logging.info(f"Found criteria {current_block} with score {score} (capped at 100 if needed)")
                else:
                    logging.warning(f"Found unknown criteria ID {current_block} in response")
                    
                # Start collecting text for this block
                current_text = [line]
            elif current_block is not None:
                # Continue adding to the current block
                current_text.append(line)
        
        # Save the last block if there is one
        if current_block is not None and current_text:
            score_blocks[current_block] = "\n".join(current_text)
            logging.info(f"Saved final block for criteria {current_block}")
        
        logging.info(f"Block-based parsing found {len(score_blocks)} criteria blocks")
        
        # If block-based parsing didn't find all scores, try line-by-line matching
        if len(score_blocks) < len(criteria_mapping):
            logging.info(f"Block parsing incomplete. Found {len(score_blocks)}/{len(criteria_mapping)} criteria. Trying line-by-line parsing.")
            for line in lines:
                line = line.strip().lower()
                if not line:
                    continue
                    
                # Try different score patterns used by Claude
                patterns = [
                    r"(?:criterion\s*)?(\d+)[:.]\s*(\d+)",  # 4: 100 or criterion 4: 100
                    r"criterion\s*#?(\d+)[:\s-]+\s*(\d+)",  # criterion #4 - 100 or criterion 4: 100
                    r"#?(\d+)\s*[:=]\s*(\d+)",             # #4 = 100 or 4 = 100
                    r"for\s+(?:criterion\s+)?#?(\d+)[,:\s]\s*(?:score is\s*)?(\d+)"  # for criterion #4, score is 100
                ]
                
                for pattern in patterns:
                    matches = re.findall(pattern, line)
                    for match in matches:
                        try:
                            criteria_num = int(match[0])
                            score = min(int(match[1]), 100)
                            
                            if criteria_num in criteria_mapping and criteria_num not in score_blocks:
                                results[criteria_num]["score"] = score
                                # Record that we found this via line matching, not block matching
                                score_blocks[criteria_num] = line
                                logging.info(f"Line parsing found criteria {criteria_num} with score {score} using pattern '{pattern}'")
                        except (ValueError, IndexError):
                            continue
        
        # Look for evidence of no education throughout the entire response
        no_education_patterns = [
            r"no education (?:credentials|mentioned|listed|found|present)",
            r"education (?:not mentioned|not found|not present|not listed|absent)",
            r"no mention of education",
            r"resume does not (?:mention|include|list|contain) (?:any )?education",
            r"couldn't find any education",
            r"no educational (?:background|information|details|qualifications)"
        ]
        
        for pattern in no_education_patterns:
            if re.search(pattern, response_text, re.IGNORECASE):
                # If any pattern matches, mark education as not found
                if 4 in results:
                    results[4]["education_found"] = False
                if 13 in results:
                    results[13]["education_found"] = False
                logging.info("Detected no education mentioned in resume")
                break
        
        # Extract explanations from the blocks
        logging.info("Extracting explanations from identified blocks")
        for criteria_id, block_text in score_blocks.items():
            if criteria_id in results:
                # The first line typically contains the score, remove it to get the explanation
                explanation_lines = block_text.split("\n")[1:]
                explanation = "\n".join(explanation_lines).strip()
                
                if explanation:
                    results[criteria_id]["explanation"] = explanation
                    # Log a truncated version of the explanation for debugging
                    short_explanation = explanation[:100] + "..." if len(explanation) > 100 else explanation
                    logging.info(f"Criteria {criteria_id}: Added explanation ({len(explanation)} chars): {short_explanation}")
                    
                    # For certain criteria, extract additional information from the explanation
                    if criteria_id == 15:  # Average Project Length
                        # Try to extract the calculated average from the explanation
                        avg_match = re.search(r"average\s+is\s+(\d+\.?\d*)", explanation, re.IGNORECASE)
                        if avg_match:
                            avg_years = float(avg_match.group(1))
                            results[criteria_id]["average_years"] = avg_years
                            logging.info(f"Extracted average project length: {avg_years} years")
                        else:
                            logging.warning(f"Could not extract average project length from explanation")
                    elif criteria_id == 4:  # Education Match
                        # Check if education was found in resume from the explanation
                        if re.search(r"no education|education not found|no mention of education", 
                                    explanation, re.IGNORECASE):
                            results[criteria_id]["education_found"] = False
                            logging.info(f"Education not found in resume per explanation")
                    elif criteria_id == 14:  # Skill Usage Duration
                        # Try to extract years of experience with the skill
                        years_match = re.search(r"(\d+\.?\d*)\s+years", explanation, re.IGNORECASE)
                        if years_match:
                            years_exp = float(years_match.group(1))
                            results[criteria_id]["years_experience"] = years_exp
                            logging.info(f"Extracted skill usage duration: {years_exp} years")
                        else:
                            logging.warning(f"Could not extract skill usage duration from explanation")
        
        # Check for missing criteria scores
        missing_criteria = [criteria_id for criteria_id in criteria_mapping.keys() if criteria_id not in score_blocks]
        if missing_criteria:
            logging.warning(f"Failed to parse scores for these criteria: {missing_criteria}")
        
        # Log the final scores for all criteria
        for criteria_id in sorted(criteria_mapping.keys()):
            if criteria_id in results:
                logging.info(f"Final score for criteria {criteria_id} ({criteria_mapping.get(criteria_id, 'unknown')}): {results[criteria_id]['score']}")
            else:
                logging.warning(f"No score for criteria {criteria_id} ({criteria_mapping.get(criteria_id, 'unknown')})")
                
        # Debug logging
        self.print_debug_info(response_text, results)
        
        logging.info("Completed parsing general analysis response")
        return results

    def print_debug_info(self, response_text: str, parsed_results: Dict) -> None:
        """
        Print debug information about parsing results

        Args:
            response_text (str): Raw response text from API
            parsed_results (Dict): Parsed results dictionary
        """
        print("\nDebug Information:")
        print("=" * 50)
        print("\nRaw Response:")
        print(response_text)
        print("\nParsed Results:")
        print(parsed_results)
        print("=" * 50)
        
    def extract_skill_duration_from_experience(self, resume_text: str, skill_name: str) -> List[Tuple[int, int, str]]:
        """
        Extract the duration periods when a skill was used, from the resume text.
        
        Args:
            resume_text (str): The resume text
            skill_name (str): The skill to look for
            
        Returns:
            List[Tuple[int, int, str]]: List of (start_year, end_year, description) tuples
        """
        # This is a placeholder implementation
        # In a real implementation, you would parse the resume to find periods when a skill was used
        return []
        
    def calculate_total_skill_duration(self, skill_periods: List[Tuple[int, int, str]]) -> int:
        """
        Calculate the total duration a skill has been used based on extracted periods.
        
        Args:
            skill_periods (List[Tuple[int, int, str]]): List of skill usage periods
            
        Returns:
            int: Total years of experience with the skill
        """
        if not skill_periods:
            return 0
            
        total_years = 0
        for start_year, end_year, _ in skill_periods:
            total_years += (end_year - start_year)
            
        return total_years
        
    def count_skill_occurrences(self, resume_text: str, skill_name: str) -> int:
        """
        Count the occurrences of a skill in the resume text.
        
        Args:
            resume_text (str): The resume text
            skill_name (str): The skill to count
            
        Returns:
            int: Number of occurrences
        """
        return resume_text.lower().count(skill_name.lower())
        
    def calculate_skill_scores(self, mentions: int, last_used_year: int, total_experience: int) -> Dict[str, int]:
        """
        Calculate various scores for a skill based on mentions, recency, and duration.
        
        Args:
            mentions (int): Number of times the skill is mentioned
            last_used_year (int): Most recent year the skill was used
            total_experience (int): Total years of experience with this skill
            
        Returns:
            Dict[str, int]: Dictionary of calculated scores
        """
        # Frequency score (5 points per mention, max 100)
        frequency_score = min(mentions * 5, 100)
        
        # Recency score (100 if used this year, -10 for each year back, min 0)
        years_since_used = self.current_year - last_used_year
        recency_score = max(100 - (years_since_used * 10), 0)
        
        # Duration score (10 points per year of experience, max 100)
        duration_score = min(total_experience * 10, 100)
        
        # Total score (average of the three)
        total_score = (frequency_score + recency_score + duration_score) // 3
        
        return {
            "frequency_score": frequency_score,
            "recency_score": recency_score,
            "duration_score": duration_score,
            "total_score": total_score
        }


class ResumeScoreUpdater:
    @staticmethod
    def update_resume_scores(db_conn: DatabaseConnection, evaluation_results: list) -> None:
        """
        Update resume scores in the database for each candidate and criteria
        
        Args:
            db_conn: DatabaseConnection instance
            evaluation_results: List of dictionaries containing evaluation results
        """
        try:
            # All skill criteria IDs that must be updated regardless of required skills
            all_skill_criteria_ids = [1, 2, 3, 9, 8]  # skill1, skill2, skill3, skill4, skill5
            
            for result in evaluation_results:
                candidate_id = result['candidate_id']
                
                # Get manager_id from shortlisted profiles
                query_manager = """
                    SELECT manager_id, job_id 
                    FROM public.job_shortlisted_profiles 
                    WHERE candidate_id = %s
                """
                db_conn.cur.execute(query_manager, (candidate_id,))
                shortlist_data = db_conn.cur.fetchone()
                
                if not shortlist_data:
                    logging.warning(f"No shortlist data found for candidate {candidate_id}")
                    continue
                    
                manager_id, job_id = shortlist_data
                
                # Always update ALL skill criteria, regardless of what's in the skills_breakdown
                for criteria_id in all_skill_criteria_ids:
                    # Get score from skills_breakdown if it exists, otherwise use 0
                    if 'skills_breakdown' in result and criteria_id in result['skills_breakdown']:
                        score_data = result['skills_breakdown'][criteria_id]
                        score = score_data.get('score', 0)
                    else:
                        # If for any reason the criteria isn't in skills_breakdown, set score to 0
                        score = 0
                    
                    # Log the score being set for debugging
                    logging.info(f"Setting score for candidate {candidate_id}, criteria {criteria_id}: {score}")
                    
                    ResumeScoreUpdater._upsert_score(
                        db_conn,
                        candidate_id=candidate_id,
                        manager_id=manager_id,
                        job_id=job_id,
                        criteria_id=criteria_id,
                        score=score
                    )
                
                # Update general scores
                if 'general_breakdown' in result:
                    for criteria_id, score_data in result['general_breakdown'].items():
                        ResumeScoreUpdater._upsert_score(
                            db_conn,
                            candidate_id=candidate_id,
                            manager_id=manager_id,
                            job_id=job_id,
                            criteria_id=criteria_id,
                            score=score_data['score']
                        )
                
            db_conn.conn.commit()
            logging.info("Successfully updated all resume scores in database")
            
        except Exception as e:
            db_conn.conn.rollback()
            logging.error(f"Error updating resume scores: {str(e)}")
            logging.error(traceback.format_exc())
            raise HTTPException(
                status_code=500,
                detail=f"Error updating resume scores: {str(e)}"
            )

    @staticmethod
    def _upsert_score(
        db_conn: DatabaseConnection,
        candidate_id: int,
        manager_id: int,
        job_id: int,
        criteria_id: int,
        score: int
    ) -> None:
        """
        Insert or update a score record in the resume_score table
        
        Args:
            db_conn: DatabaseConnection instance
            candidate_id: ID of the candidate
            manager_id: ID of the manager
            job_id: ID of the job
            criteria_id: ID of the scoring criteria
            score: The calculated score
        """
        try:
            # Ensure score is an integer
            try:
                score = int(score)
            except (TypeError, ValueError):
                score = 0  # Default to 0 if score is not a valid number
            
            # Check if record exists
            query_check = """
                SELECT score_id 
                FROM public.resume_scores
                WHERE candidate_id = %s 
                AND manager_id = %s 
                AND job_id = %s 
                AND criteria_id = %s
                AND is_deleted = false
            """
            db_conn.cur.execute(query_check, (candidate_id, manager_id, job_id, criteria_id))
            existing_record = db_conn.cur.fetchone()
            
            if existing_record:
                # Update existing record
                query_update = """
                    UPDATE public.resume_scores 
                    SET score = %s,
                        updated_at = CURRENT_TIMESTAMP
                    WHERE score_id = %s
                """
                db_conn.cur.execute(query_update, (score, existing_record[0]))
            else:
                # Insert new record
                query_insert = """
                    INSERT INTO public.resume_scores 
                    (candidate_id, manager_id, job_id, criteria_id, score, created_on, updated_at, is_deleted)
                    VALUES (%s, %s, %s, %s, %s, CURRENT_TIMESTAMP, CURRENT_TIMESTAMP, false)
                """
                db_conn.cur.execute(query_insert, (candidate_id, manager_id, job_id, criteria_id, score))
                
        except Exception as e:
            logging.error(f"Error upserting score for candidate {candidate_id}, criteria {criteria_id}: {str(e)}")
            raise HTTPException(
                status_code=500,
                detail=f"Error upserting score record: {str(e)}"
            )

def evaluate_candidates_comprehensive(job_id: int):
    """
    Evaluate all shortlisted candidates and update their scores in the database
    using database-driven scoring configuration.

    Args:
        job_id (int): The ID of the job to evaluate candidates for

    Returns:
        List[Dict]: List of final scores for all candidates
    """
    # Get API key from environment
    api_key = os.getenv("CLAUDE_API_KEY")
    if not api_key:
        raise HTTPException(status_code=500, detail="CLAUDE_API_KEY not found in environment variables")
    
    # Set up logging if not already configured
    if not logging.getLogger().handlers:
        logging.basicConfig(
            level=logging.INFO,
            format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
        )
    
    # Check for AWS credentials
    aws_access_key_id = os.getenv("AWS_ACCESS_KEY_ID")
    aws_secret_access_key = os.getenv("AWS_SECRET_ACCESS_KEY")
    
    if not aws_access_key_id or not aws_secret_access_key:
        logging.error("AWS credentials not found in environment variables")
        raise HTTPException(
            status_code=500, 
            detail="AWS credentials not found. Please set AWS_ACCESS_KEY_ID and AWS_SECRET_ACCESS_KEY environment variables."
        )
    
    try:
        with DatabaseConnection() as db_conn:
            # Get job description
            try:
                job_desc = JobDataAccess.get_job_description(job_id)
                job_years_exp = JobDataAccess.get_years_of_experience(job_id)
                
                if not job_desc:
                    raise HTTPException(
                        status_code=404,
                        detail=f"Job description not found for job_id {job_id}",
                    )
            except HTTPException as he:
                # Re-raise HTTPException to preserve status code and detail
                raise he
            except Exception as e:
                logging.error(f"Error retrieving job description: {str(e)}")
                logging.error(traceback.format_exc())
                raise HTTPException(
                    status_code=500, 
                    detail=f"Error retrieving job description: {str(e)}"
                )

            # Get shortlisted profiles
            try:
                shortlisted_profiles = ShortlistDataAccess.get_shortlisted_profiles(
                    db_conn, job_id
                )
                logging.info(f"Found {len(shortlisted_profiles)} shortlisted profiles for job_id {job_id}")
                if not shortlisted_profiles:
                    raise HTTPException(
                        status_code=404,
                        detail=f"No shortlisted candidates found for job_id {job_id}. Please ensure candidates have been shortlisted for this job.",
                    )
            except HTTPException as he:
                # Re-raise HTTPException to preserve status code and detail
                raise he
            except Exception as e:
                logging.error(f"Error retrieving shortlisted profiles: {str(e)}")
                logging.error(traceback.format_exc())
                raise HTTPException(
                    status_code=500, 
                    detail=f"Error retrieving shortlisted profiles: {str(e)}"
                )

            # Initialize analyzers
            try:
                skills_fetcher = SkillsFetcher()
                skills_fetcher.fetch_skills(job_id)
                required_skills = skills_fetcher.get_skills_dict()
                logging.info(f"Found {len(required_skills)} required skills for job_id {job_id}")

                if not required_skills:
                    logging.warning(f"No required skills found for job_id {job_id}")
                    raise HTTPException(
                        status_code=404, 
                        detail=f"No required skills found for job_id {job_id}"
                    )
            except HTTPException as he:
                raise he
            except Exception as e:
                logging.error(f"Error fetching required skills: {str(e)}")
                logging.error(traceback.format_exc())
                raise HTTPException(
                    status_code=500, 
                    detail=f"Error fetching required skills: {str(e)}"
                )

            # Initialize ResumeAnalyzer and load scoring configurations from database
            resume_analyzer = ResumeAnalyzer(api_key)
            resume_analyzer.load_scoring_configs(db_conn)

            # Create/check resume directory
            resume_directory = "RESUME"
            if not os.path.exists(resume_directory):
                os.makedirs(resume_directory)
                logging.info(f"Created resume directory: {resume_directory}")

            # Process each candidate
            evaluation_results = []
            final_scores = []
            errors = []

            for profile in shortlisted_profiles:
                candidate_id = profile["candidate_id"]
                logging.info(f"Processing candidate ID: {candidate_id}")
                local_resume_path = None

                try:
                    # Get and process resume
                    candidates = CandidateDataAccess.get_candidate_details(
                        db_conn, candidate_id
                    )
                    if not candidates or not candidates[0].get("s3_key_resume"):
                        error_msg = f"No resume found for candidate_id {candidate_id}"
                        logging.warning(error_msg)
                        errors.append(error_msg)
                        continue

                    s3_key_resume = candidates[0]["s3_key_resume"]
                    local_resume_path = os.path.join(
                        resume_directory,
                        f"{candidate_id}_resume{os.path.splitext(s3_key_resume)[1]}",
                    )

                    # Download and read resume
                    try:
                        logging.info(f"Downloading resume from S3: {s3_key_resume}")
                        S3FileDownload.download_s3_file(s3_key_resume, local_resume_path)
                    except Exception as e:
                        error_msg = f"Error downloading resume for candidate {candidate_id}: {str(e)}"
                        logging.error(error_msg)
                        logging.error(traceback.format_exc())
                        errors.append(error_msg)
                        continue

                    resume_text = DocumentParser.read_file(local_resume_path)
                    if not resume_text:
                        error_msg = f"Could not extract text from resume for candidate_id {candidate_id}"
                        logging.warning(error_msg)
                        errors.append(error_msg)
                        continue

                    # Skills Analysis with database-driven scoring
                    try:
                        skills_scores = resume_analyzer.analyze_resume_skills(
                            resume_text, required_skills
                        )
                        
                        if not isinstance(skills_scores, dict):
                            error_msg = f"Skills analysis failed for candidate_id {candidate_id}: {skills_scores}"
                            logging.warning(error_msg)
                            errors.append(error_msg)
                            continue
                            
                        # Calculate average skill score using only non-zero scores
                        non_zero_scores = [s["score"] for s in skills_scores.values() if s["score"] > 0]
                        if non_zero_scores:
                            skills_avg = sum(non_zero_scores) / len(non_zero_scores)
                        else:
                            skills_avg = 0
                            
                        logging.info(f"Completed skills analysis for candidate {candidate_id} with avg score: {skills_avg}")
                    except Exception as e:
                        error_msg = f"Error in skills analysis for candidate {candidate_id}: {str(e)}"
                        logging.error(error_msg)
                        logging.error(traceback.format_exc())
                        errors.append(error_msg)
                        continue

                    # General Analysis
                    try:
                        general_analysis = resume_analyzer.general_analyze_resume(
                            resume_text, job_desc, job_id, job_years_exp
                        )
                        
                        if not isinstance(general_analysis, dict):
                            error_msg = f"General analysis failed for candidate_id {candidate_id}: {general_analysis}"
                            logging.warning(error_msg)
                            errors.append(error_msg)
                            continue
                            
                        detailed_scores = general_analysis.get("detailed_scores", {})
                        general_avg = general_analysis.get("overall_score", 0)
                        logging.info(f"Completed general analysis for candidate {candidate_id} with avg score: {general_avg}")
                    except Exception as e:
                        error_msg = f"Error in general analysis for candidate {candidate_id}: {str(e)}"
                        logging.error(error_msg)
                        logging.error(traceback.format_exc())
                        errors.append(error_msg)
                        continue

                    # Calculate final score
                    final_score = (skills_avg + general_avg) / 2

                    # Store evaluation results
                    result = {
                        "candidate_id": candidate_id,
                        "final_score": final_score,
                        "skills_score": skills_avg,
                        "general_score": general_avg,
                        "skills_breakdown": skills_scores,
                        "general_breakdown": detailed_scores,
                        "manager_id": profile["manager_id"],
                    }
                    evaluation_results.append(result)
                    logging.info(f"Evaluation complete for candidate {candidate_id} with final score: {final_score}")

                    # Add to final scores list
                    final_scores.append(
                        {"candidate_id": candidate_id, "final_score": final_score}
                    )

                except Exception as e:
                    error_msg = f"Error processing candidate {candidate_id}: {str(e)}"
                    logging.error(error_msg)
                    logging.error(traceback.format_exc())
                    errors.append(error_msg)
                    continue
                finally:
                    # Cleanup downloaded resume
                    if local_resume_path and os.path.exists(local_resume_path):
                        try:
                            os.remove(local_resume_path)
                            logging.info(f"Removed temporary file: {local_resume_path}")
                        except Exception as e:
                            logging.warning(f"Failed to remove temporary file {local_resume_path}: {str(e)}")

            # Update scores in database
            if evaluation_results:
                try:
                    logging.info(f"Updating database with scores for {len(evaluation_results)} candidates")
                    ResumeScoreUpdater.update_resume_scores(db_conn, evaluation_results)
                except Exception as e:
                    error_msg = f"Error updating scores in database: {str(e)}"
                    logging.error(error_msg)
                    logging.error(traceback.format_exc())
                    errors.append(error_msg)
                    # Continue despite database update errors

            if not final_scores:
                if errors:
                    error_summary = "; ".join(errors[:5])
                    if len(errors) > 5:
                        error_summary += f"; and {len(errors) - 5} more errors"
                    logging.error(f"No candidates were successfully evaluated: {error_summary}")
                    raise HTTPException(
                        status_code=500,
                        detail=f"Failed to evaluate any candidates: {error_summary}"
                    )
                else:
                    logging.error(f"No candidates were successfully evaluated for job_id {job_id}")
                    raise HTTPException(
                        status_code=404,
                        detail=f"No candidates were successfully evaluated for job_id {job_id}"
                    )

            logging.info(f"Successfully evaluated {len(final_scores)} candidates for job_id {job_id}")
            return final_scores

    except HTTPException as he:
        # Re-raise HTTPExceptions to preserve their status code and detail
        raise he
    except Exception as e:
        # Convert generic exceptions to HTTPExceptions with detail
        error_message = str(e)
        logging.error(f"Evaluation failed: {error_message}")
        logging.error(traceback.format_exc())
        raise HTTPException(status_code=500, detail=error_message)
